from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from plagiarismchecker.algorithm import main
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
from plagiarismchecker.algorithm import resumep
import PyPDF2 
from django.views.decorators.csrf import csrf_exempt

# Create your views here.
#home
def home(request):
    return render(request, 'pc/index.html') 
@csrf_exempt
def test(request):
    text=request.GET.get('text')
    print("text",text)
    
    print("request is welcome text")
    if text: 
        percent,link = main.findSimilarity(text)
        print("================555====================")
        percent = round(percent,2)
    print("===================LINK=================")
    print(len(link),percent)
    print("=================Percent===================")
    data = [{'link': link, 'percentage': percent}]
    return JsonResponse(data, safe=False)
    #return HttpResponse({'link': link, 'percent': percent})

@csrf_exempt
def resume(request):
    link=request.GET.get('resume')
    print("link",link)
    print("====================================")
    resume_score=resumep.resume_parser(link)
    print("resume",resume_score)
    print("====================================")
    return HttpResponse(resume_score)
    #return render(request, 'pc/resume.html')

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extracting text from page 
        print(pageObj.extractText()) 

        # closing the pdf file object 
        pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#text compare
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result = fileSimilarity.findFileSimilarity(value1,value2)
    
    print("Output..................!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
